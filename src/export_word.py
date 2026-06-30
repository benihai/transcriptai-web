"""ייצוא סיכום פגישה לקובץ Word (.docx) עם עיצוב RTL בעברית."""
from __future__ import annotations

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ACCENT = RGBColor(0x6c, 0x5c, 0xe7)


def _rtl(paragraph):
    """הופך פסקה לכיוון ימין-לשמאל (כולל ה-runs שלה)."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = paragraph._p.get_or_add_pPr()
    if pPr.find(qn("w:bidi")) is None:
        pPr.append(OxmlElement("w:bidi"))
    for run in paragraph.runs:
        rPr = run._r.get_or_add_rPr()
        rPr.append(OxmlElement("w:rtl"))


def _para(doc, text, *, size=11, bold=False, color=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.name = "David"
    if color:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    _rtl(p)
    return p


def _heading(doc, text):
    return _para(doc, text, size=14, bold=True, color=ACCENT, space_after=4)


def _bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.name = "David"
    _rtl(p)
    return p


def build_docx(meeting: dict, out_path: str) -> str:
    """בונה מסמך Word מתוך נתוני פגישה ושומר ב-out_path."""
    s = meeting.get("summary", {})
    doc = Document()

    # כותרת
    _para(doc, s.get("title", "סיכום פגישה"), size=20, bold=True, space_after=2)
    if meeting.get("date"):
        _para(doc, meeting["date"], size=10, color=RGBColor(0x88, 0x88, 0x88), space_after=12)

    # תקציר
    if s.get("summary"):
        _heading(doc, "תקציר")
        _para(doc, s["summary"], space_after=12)

    # פירוט לפי נושאים
    if s.get("topics"):
        _heading(doc, "פירוט הנושאים")
        for t in s["topics"]:
            if t.get("title"):
                _para(doc, t["title"], size=12, bold=True, space_after=2)
            for pt in t.get("points", []):
                _bullet(doc, pt)
            doc.add_paragraph()

    # נקודות עיקריות
    if s.get("key_points"):
        _heading(doc, "נקודות עיקריות")
        for p in s["key_points"]:
            _bullet(doc, p)
        doc.add_paragraph()

    # החלטות
    if s.get("decisions"):
        _heading(doc, "החלטות")
        for d in s["decisions"]:
            _bullet(doc, d)
        doc.add_paragraph()

    # משימות לביצוע
    if s.get("action_items"):
        _heading(doc, "משימות לביצוע")
        for a in s["action_items"]:
            parts = [f"☐ {a.get('task','')}"]
            extra = []
            if a.get("owner"):
                extra.append(f"אחראי: {a['owner']}")
            if a.get("due"):
                extra.append(f"יעד: {a['due']}")
            line = parts[0] + (f"  ({', '.join(extra)})" if extra else "")
            _para(doc, line, space_after=4)
        doc.add_paragraph()

    # תמלול מלא
    if meeting.get("transcript"):
        _heading(doc, "תמלול מלא")
        _para(doc, meeting["transcript"], size=10,
              color=RGBColor(0x55, 0x55, 0x55))

    doc.save(out_path)
    return out_path
